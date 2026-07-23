#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
v4.2 Weekly Momentum Backtest (fixed version)
- 修复: 默认参数改为v4.2 (LB=5, max_dev=10, top_n=3)
- 新增: 沪深300市场状态过滤 (三周动量 > -1%)
- 修复: G3过滤动量计算 (使用wk[-4]而非wk[-3])
- 数据源: 使用本地历史文件 (与扫描一致)

Usage: python backtest_v4_fixed.py [--max-dev 10] [--top-n 3] [--lb 5]
"""

import json, os, sys, glob, statistics
from datetime import datetime
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# HISTORY_DIR = r'D:\QClaw_Trading\data\history_long'  # ⚠️ OVERWRITTEN on 2026-06-13, only 300 records
# HISTORY_DIR = r'D:\QClaw_Trading\data\history_long'  # 旧数据, 过期
HISTORY_DIR = r'D:\QClaw_Trading\data\history_long_v2'  # ✅ AKShare腾讯前复权(194 ETF, 2010+)
POOL_FILE = r'D:\QClaw_Trading\data\etf_pool_V1_full.json'
OUTPUT_DIR = r'D:\QClaw_Trading\backtest_results'


def load_pool():
    with open(POOL_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data.get('data', data.get('etfs', []))


def load_history(code):
    """Load history from local file (consistent with scan)."""
    for pat in [code, f'sh{code}', f'sz{code}']:
        matches = glob.glob(os.path.join(HISTORY_DIR, f'{pat}.json'))
        if not matches:
            matches = glob.glob(os.path.join(HISTORY_DIR, f'*{code}.json'))
        if matches:
            try:
                with open(matches[0], 'r', encoding='utf-8') as f:
                    raw = f.read().replace('NaN', 'null')
                d = json.loads(raw)
                recs = d.get('records', []) if isinstance(d, dict) else d
                weeks = {}
                for r in recs:
                    if isinstance(r, dict):
                        ds, cl = r.get('date', ''), float(r.get('close', 0))
                    else:
                        ds, cl = str(r[0]), float(r[2])
                    try:
                        dt = datetime.strptime(ds, '%Y-%m-%d')
                        y, w, _ = dt.isocalendar()
                        week_key = f'{y}-W{w:02d}'
                        # Keep latest close for the week
                        if week_key not in weeks or ds > weeks[week_key][0]:
                            weeks[week_key] = (ds, cl)
                    except:
                        pass
                # Return sorted list of (week, close)
                sw = sorted(weeks.items())
                return [(w, cl) for w, (ds, cl) in sw]
            except:
                continue
    return None


def load_hs300_momentum():
    """Load HS300 3-week momentum for market filter."""
    hs300 = load_history('000300')
    if not hs300 or len(hs300) < 5:
        return None
    
    weeks = [w for w, c in hs300]
    closes = [c for w, c in hs300]
    
    # Calculate 3-week momentum for each week
    mom_map = {}
    for i in range(5, len(closes)):  # Need at least 5 weeks for MA21
        w = weeks[i]
        mom = closes[i] / closes[i-5] - 1  # 3-week momentum (actually 4-week span)
        mom_map[w] = mom
    
    return mom_map


def export_to_word(args, label, total_ret, ann_ret, max_dd, sharpe, calmar, win_rate,
                  n, years, n_buys, n_sells, eq_curve, yg, output_dir, ts):
    """Export backtest results to Word .docx file."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return None

    doc = Document()

    # Title
    t = doc.add_heading('周线动量策略 回测报告 (v4.2)', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'策略参数：{label}')
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()

    # Parameters table
    doc.add_heading('一、回测参数', 1)
    pt = doc.add_table(rows=1, cols=2)
    pt.style = 'Light Grid Accent 1'
    ph = pt.rows[0].cells
    ph[0].text = '参数'
    ph[1].text = '值'
    for k, v in vars(args).items():
        row = pt.add_row()
        row.cells[0].text = str(k)
        row.cells[1].text = str(v)
    doc.add_paragraph()

    # Stats
    doc.add_heading('二、回测结果', 1)
    doc.add_paragraph(f'回测区间：{args.start} ~ {args.end}（{n}周，{years:.1f}年）')
    p = doc.add_paragraph()
    p.add_run(f'总收益率：{total_ret:+.1f}%').bold = True
    p = doc.add_paragraph()
    p.add_run(f'年化收益率：{ann_ret:+.1f}%').bold = True
    p = doc.add_paragraph()
    p.add_run(f'最大回撤：{max_dd*100:.1f}%').bold = True
    p = doc.add_paragraph()
    p.add_run(f'夏普比率：{sharpe:.2f}').bold = True
    p = doc.add_paragraph()
    p.add_run(f'Calmar比率：{calmar:.2f}').bold = True
    doc.add_paragraph(f'周胜率：{win_rate:.1f}%')
    doc.add_paragraph(f'交易次数：{n_buys}次买入 / {n_sells}次卖出')
    doc.add_paragraph()

    # Yearly breakdown
    doc.add_heading('三、逐年表现', 1)
    yt = doc.add_table(rows=1, cols=4)
    yt.style = 'Light Grid Accent 1'
    yh = yt.rows[0].cells
    yh[0].text = '年份'
    yh[1].text = '收益率'
    yh[2].text = '最大回撤'
    yh[3].text = '平均持仓占比'
    for yr in sorted(yg):
        es = [e['eq'] for e in yg[yr]]
        if not es or es[0] <= 0:
            continue
        ret = (es[-1] / es[0] - 1) * 100
        # 正确计算年内最大回撤（峰值必须在该周之前）
        pk = es[0]
        max_dd = 0
        for eq in es:
            if eq > pk:
                pk = eq
            dd = (eq / pk - 1) * 100
            if dd < max_dd:
                max_dd = dd
        dd = max_dd
        hold_pct = statistics.mean([e['nh'] for e in yg[yr]]) / args.top_n * 100
        row = yt.add_row()
        row.cells[0].text = yr
        row.cells[1].text = f'{ret:+.1f}%'
        row.cells[2].text = f'{dd:.1f}%'
        row.cells[3].text = f'{hold_pct:.0f}%'
    doc.add_paragraph()

    # Equity curve summary
    doc.add_heading('四、权益曲线（摘要）', 1)
    et = doc.add_table(rows=1, cols=4)
    et.style = 'Light Grid Accent 1'
    eh = et.rows[0].cells
    eh[0].text = '周次'
    eh[1].text = '权益'
    eh[2].text = '持仓数'
    eh[3].text = '占比'
    step = max(1, len(eq_curve) // 30)
    for i, e in enumerate(eq_curve):
        if i % step == 0 or i == len(eq_curve) - 1:
            row = et.add_row()
            row.cells[0].text = e['w']
            row.cells[1].text = f'{e["eq"]:.4f}'
            row.cells[2].text = str(e['nh'])
            row.cells[3].text = f'{e["nh"]/args.top_n*100:.0f}%'

    # Save
    fname = f'bt_v4_2_{args.ma_s}_{args.ma_l}_{args.lb}_{int(args.max_dev)}_{args.top_n}_{ts}.docx'
    fp = os.path.join(output_dir, fname)
    try:
        doc.save(fp)
        return fp
    except Exception:
        return None


def main():
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('--max-dev', type=float, default=10)  # v4.3: 10%
    ap.add_argument('--top-n', type=int, default=2)      # v4.3: 2只
    ap.add_argument('--lb', type=int, default=3)          # v4.3: LB=3
    ap.add_argument('--ma-s', type=int, default=5)
    ap.add_argument('--ma-l', type=int, default=21)
    ap.add_argument('--capital', type=float, default=1.0)
    ap.add_argument('--start', type=str, default='2010-W01')
    ap.add_argument('--end', type=str, default='2026-W18')
    ap.add_argument('--hs300-threshold', type=float, default=-100.0)  # HS300过滤阈值(≤-100禁用)
    ap.add_argument('--output', type=str, default=None)  # 指定输出文件名
    ap.add_argument('--dynamic-weight', action='store_true', help='启用动态仓位管理（按动量强度分配权重）')
    ap.add_argument('--volume-filter', action='store_true', help='启用成交量过滤（本周成交量>20周均量*1.2）')
    args = ap.parse_args()

    label = f"MA{args.ma_s}/{args.ma_l} LB{args.lb} D{args.max_dev} H{args.top_n}"
    print(f"{'='*60}")
    print(f"  v4.2 Backtest (fixed): {label}")
    print(f"  Range: {args.start} ~ {args.end}")
    print(f"{'='*60}\n")

    etfs = load_pool()
    print("Loading...")

    all_series = {}
    code_info = {}
    missing = 0
    for etf in etfs:
        code = etf['code']
        s = load_history(code)
        if s and len(s) >= 30:
            all_series[code] = s
            code_info[code] = {'name': etf.get('name', ''), 'cat': etf.get('category', '')}
        else:
            missing += 1

    print(f"  Loaded: {len(all_series)}/{len(etfs)}, missing: {missing}")

    weeks_set = set()
    for s in all_series.values():
        for w, c in s:
            weeks_set.add(w)
    all_weeks = sorted(w for w in weeks_set if args.start <= w <= args.end)
    print(f"  Weeks: {len(all_weeks)} ({all_weeks[0]} ~ {all_weeks[-1]})\n")

    # Load HS300 momentum for market filter
    hs300_mom = load_hs300_momentum()
    if hs300_mom:
        print(f"  HS300 momentum loaded: {len(hs300_mom)} weeks")
    else:
        print(f"  Warning: HS300 data not available, market filter disabled")

    def closes_until(code, week):
        return [c for w, c in all_series.get(code, []) if w <= week]

    def close_at(code, week):
        for w, c in all_series.get(code, []):
            if w == week: return c
            if w > week: return None
        return None

    # Signal function
    def get_signal(code, week):
        """Return (close, mom) if passes filter, else None."""
        cs = closes_until(code, week)
        n = len(cs)
        if n < args.ma_l + 1:
            return None
        
        # Market filter: HS300 3-week momentum > threshold
        if hs300_mom and week in hs300_mom:
            if hs300_mom[week] <= args.hs300_threshold / 100:
                return None
        
        price = cs[-1]
        ma_s = sum(cs[-args.ma_s:]) / args.ma_s
        ma_l = sum(cs[-args.ma_l:]) / args.ma_l
        mom = cs[-1] / cs[-args.lb] - 1 if n > args.lb else None
        dev = price / ma_l - 1
        
        if mom is None or mom <= 0:
            return None
        if not (price > ma_s > ma_l):
            return None
        if dev > args.max_dev / 100.0:
            return None
        
        # G3 filter: 3w >= 0% AND 1w >= -1%
        g3_pass = True
        if len(cs) >= 2:
            mom1w = cs[-1] / cs[-2] - 1
            if mom1w < -0.01:
                g3_pass = False
        if len(cs) >= 4:  # FIXED: Need 4 weeks for 3-week momentum (cs[-4])
            mom3w = cs[-1] / cs[-4] - 1
            if mom3w < 0:
                g3_pass = False
        
        if not g3_pass:
            return None
        
        return {'code': code, 'close': price, 'mom': mom, 'dev': dev}

    # Backtest: signal at T, execute at T+1
    portfolio = {}  # code -> {weight, buy_price, hwm}
    cash = args.capital
    eq_curve = []
    trades = []
    n_buys = n_sells = 0

    # We need T+1 week to execute, so iterate signal weeks
    # Process pairs: (signal_week, exec_week)
    for i in range(len(all_weeks) - 1):
        sig_week = all_weeks[i]
        exec_week = all_weeks[i + 1]

        # === 1. Signal at sig_week ===
        candidates = []
        for code in all_series:
            sig = get_signal(code, sig_week)
            if sig:
                candidates.append(sig)

        candidates.sort(key=lambda x: x['mom'], reverse=True)

        # Category dedup
        cats = set()
        target = []
        for c in candidates:
            cat = code_info.get(c['code'], {}).get('cat', '') or c['code']
            if cat not in cats:
                cats.add(cat)
                target.append(c)
        target = target[:args.top_n]
        target_codes = {t['code'] for t in target}

        # === 2. Check stop losses on current holdings at exec_week prices ===
        # (stop loss is checked each week, not just on rebalance)
        # But for simplicity, we check at exec_week

        # === 3. Execute at exec_week ===
        # First: mark-to-market all holdings at exec_week
        for code in portfolio:
            p = close_at(code, exec_week)
            if p and p > portfolio[code]['hwm']:
                portfolio[code]['hwm'] = p

        # Sells
        to_sell = []
        for code, pos in list(portfolio.items()):
            p = close_at(code, exec_week)
            if p is None:
                to_sell.append((code, 'no_data'))
                continue
            cost_pnl = p / pos['buy_price'] - 1
            hwm_pnl = p / pos['hwm'] - 1
            if cost_pnl <= -0.08 or hwm_pnl <= -0.10:
                to_sell.append((code, 'stop'))
            elif code not in target_codes:
                to_sell.append((code, 'rebalance'))

        for code, reason in to_sell:
            pos = portfolio.pop(code)
            p = close_at(code, exec_week) or pos['buy_price']
            cash += pos['weight'] * p
            pnl = (p / pos['buy_price'] - 1) * 100
            trades.append({'w': exec_week, 'act': 'S', 'code': code,
                          'pnl': round(pnl, 2), 'reason': reason})
            n_sells += 1

        # Equity after sells
        equity = cash + sum(pos['weight'] * (close_at(c, exec_week) or pos['buy_price'])
                          for c, pos in portfolio.items())

        # Buys
        slots = args.top_n - len(portfolio)
        if slots > 0 and equity > 0:
            buy_list = [t for t in target if t['code'] not in portfolio]
            slot_val = equity / args.top_n
            for bc in buy_list[:slots]:
                exec_price = close_at(bc['code'], exec_week)
                if exec_price is None or exec_price <= 0:
                    continue
                weight = slot_val / exec_price
                cost = weight * exec_price
                if cost > cash * 0.98:
                    weight = cash * 0.98 / exec_price
                    cost = weight * exec_price
                if weight <= 0:
                    break
                cash -= cost
                portfolio[bc['code']] = {'weight': weight, 'buy_price': exec_price, 'hwm': exec_price}
                trades.append({'w': exec_week, 'act': 'B', 'code': bc['code'],
                              'price': round(exec_price, 4)})
                n_buys += 1

        # Record equity
        equity = cash + sum(pos['weight'] * (close_at(c, exec_week) or pos['buy_price'])
                          for c, pos in portfolio.items())
        eq_curve.append({'w': exec_week, 'eq': equity, 'nh': len(portfolio),
                        'cash': round(cash, 4), 'holds': list(portfolio.keys())})

    # === Stats ===
    eqs = [e['eq'] for e in eq_curve]
    n = len(eqs)
    if n < 2:
        print("Not enough data"); return

    init, final = eqs[0], eqs[-1]
    total_ret = (final / init - 1) * 100
    years = n / 52
    ann_ret = ((final / init) ** (1 / years) - 1) * 100 if years > 0 else 0

    peak = eqs[0]
    max_dd = 0
    for eq in eqs:
        if eq > peak: peak = eq
        dd = eq / peak - 1
        if dd < max_dd: max_dd = dd

    w_rets = [eqs[i] / eqs[i-1] - 1 for i in range(1, n) if eqs[i-1] > 0]
    if w_rets:
        avg_w = statistics.mean(w_rets)
        std_w = statistics.stdev(w_rets) if len(w_rets) > 1 else 1e-9
        sharpe = (avg_w * 52 - 0.02) / (std_w * 52**0.5)
        calmar = ann_ret / abs(max_dd * 100) if max_dd else 0
        win_rate = sum(1 for r in w_rets if r > 0) / len(w_rets) * 100
    else:
        sharpe = calmar = win_rate = 0

    print(f"{'='*60}")
    print(f"  RESULTS: {label}")
    print(f"{'='*60}\n")
    print(f"  Period:     {all_weeks[0]} ~ {all_weeks[-1]} ({n}w, {years:.1f}y)")
    print(f"  Total Ret:  {total_ret:+.1f}%")
    print(f"  Annual:     {ann_ret:+.1f}%")
    print(f"  Max DD:     {max_dd*100:.1f}%")
    print(f"  Sharpe:     {sharpe:.2f}")
    print(f"  Calmar:     {calmar:.2f}")
    print(f"  Win Rate:   {win_rate:.1f}%")
    print(f"  Trades:     {n_buys}B / {n_sells}S")

    # Yearly
    print(f"\n  {'Year':<6} {'Ret':>7} {'DD':>7} {'Hold%':>7}")
    print(f"  {'-'*30}")
    yg = defaultdict(list)
    for e in eq_curve:
        yg[e['w'][:4]].append(e)
    for yr in sorted(yg):
        es = [e['eq'] for e in yg[yr]]
        if es[0] <= 0: continue
        ret = (es[-1] / es[0] - 1) * 100
        # 正确计算年内最大回撤（峰值必须在该周之前）
        pk = es[0]
        yearly_max_dd = 0
        for eq in es:
            if eq > pk:
                pk = eq
            dd = (eq / pk - 1) * 100
            if dd < yearly_max_dd:
                yearly_max_dd = dd
        dd = yearly_max_dd
        hold_pct = statistics.mean(e['nh'] for e in yg[yr]) / args.top_n * 100
        print(f"  {yr:<6} {ret:>+6.1f}% {dd:>6.1f}% {hold_pct:>6.0f}%")

    # Save JSON
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    if args.output:
        fp = args.output if os.path.isabs(args.output) else os.path.join(OUTPUT_DIR, args.output)
    else:
        fname = f'bt_v4_2_{args.ma_s}_{args.ma_l}_{args.lb}_{int(args.max_dev)}_{args.top_n}_{ts}.json'
        fp = os.path.join(OUTPUT_DIR, fname)
    with open(fp, 'w', encoding='utf-8') as f:
        json.dump({
            'label': label, 'ts': datetime.now().isoformat(),
            'params': vars(args),
            'stats': {'total_ret': total_ret, 'ann_ret': ann_ret, 'max_dd': max_dd,
                      'sharpe': sharpe, 'calmar': calmar, 'win_rate': win_rate,
                      'n': n, 'years': years, 'n_buys': n_buys, 'n_sells': n_sells},
            'equity': [{'w': e['w'], 'eq': round(e['eq'], 6), 'nh': e['nh']} for e in eq_curve],
        }, f, ensure_ascii=False, indent=2)
    print(f"\nSaved: {fp}")

    # Export to Word
    if HAS_DOCX:
        doc_path = export_to_word(
            args, label, total_ret, ann_ret, max_dd, sharpe, calmar, win_rate,
            n, years, n_buys, n_sells, eq_curve, yg, OUTPUT_DIR, ts
        )
        if doc_path:
            print(f"Word saved: {doc_path}")
    else:
        print("Tip: pip install python-docx to enable Word export")


if __name__ == '__main__':
    main()
